from __future__ import annotations

import re

import httpx

from app.connectors.google_drive.auth import read_access_token
from app.connectors.google_drive.constants import (
    DOCX_MIME,
    EXPORT_MIME_TYPES,
    GOOGLE_DRIVE_DOWNLOAD_URL,
    GOOGLE_DRIVE_EXPORT_URL,
    XLSX_MIME,
)
from app.connectors.plugin_types import ConnectorFetchInput, ConnectorFetchedDocument

MAX_CONTENT_BYTES = 2 * 1024 * 1024


async def fetch_google_document(input: ConnectorFetchInput) -> ConnectorFetchedDocument:
    access_token = read_access_token(input.credentials)
    mime_type = input.ref.content_type or ""

    async with httpx.AsyncClient() as client:
        if mime_type in EXPORT_MIME_TYPES:
            url = GOOGLE_DRIVE_EXPORT_URL.replace("{file_id}", input.ref.external_id)
            response = await client.get(
                url,
                params={"mimeType": EXPORT_MIME_TYPES[mime_type]},
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=60.0,
            )
        else:
            url = GOOGLE_DRIVE_DOWNLOAD_URL.replace("{file_id}", input.ref.external_id)
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=60.0,
            )

    if response.status_code == 404:
        return ConnectorFetchedDocument(
            content=None,
            content_type=mime_type,
            metadata={**input.ref.metadata, "fetch_status": 404},
        )

    if response.status_code != 200:
        from app.types.sync_errors import SyncErrorCode, SyncPipelineError
        raise SyncPipelineError(
            code=SyncErrorCode.FETCH_FAILED,
            stage="fetch",
            message=response.text,
            retriable=response.status_code >= 500,
        )

    raw = response.content[:MAX_CONTENT_BYTES]

    if mime_type == "application/pdf":
        try:
            import io
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
            return ConnectorFetchedDocument(
                content=text or None,
                content_type=mime_type,
                metadata=input.ref.metadata,
            )
        except Exception as exc:
            from app.types.sync_errors import SyncErrorCode, SyncPipelineError
            raise SyncPipelineError(
                code=SyncErrorCode.PARSE_FAILED,
                stage="normalize",
                message="Failed to parse PDF",
                retriable=False,
            ) from exc

    if mime_type == DOCX_MIME:
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ConnectorFetchedDocument(
                content=text or None,
                content_type=mime_type,
                metadata=input.ref.metadata,
            )
        except Exception as exc:
            from app.types.sync_errors import SyncErrorCode, SyncPipelineError
            raise SyncPipelineError(
                code=SyncErrorCode.PARSE_FAILED,
                stage="normalize",
                message="Failed to parse DOCX",
                retriable=False,
            ) from exc

    if mime_type == XLSX_MIME:
        try:
            import io
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append("\t".join(cells))
            text = "\n".join(lines)
            return ConnectorFetchedDocument(
                content=text or None,
                content_type=mime_type,
                metadata=input.ref.metadata,
            )
        except Exception as exc:
            from app.types.sync_errors import SyncErrorCode, SyncPipelineError
            raise SyncPipelineError(
                code=SyncErrorCode.PARSE_FAILED,
                stage="normalize",
                message="Failed to parse XLSX",
                retriable=False,
            ) from exc

    text = raw.decode("utf-8", errors="replace")
    if not text.strip():
        return ConnectorFetchedDocument(
            content=None,
            content_type=mime_type,
            metadata=input.ref.metadata,
        )

    if mime_type == "text/html":
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s{2,}", " ", text).strip() or None  # type: ignore[assignment]
        return ConnectorFetchedDocument(
            content=text,
            content_type=mime_type,
            metadata=input.ref.metadata,
        )

    return ConnectorFetchedDocument(
        content=text,
        content_type=mime_type,
        metadata=input.ref.metadata,
    )
